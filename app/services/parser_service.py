"""
文档解析服务

支持多种格式文档的文本提取和分段处理
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

import pypdf
from docx import Document
import markdown
from loguru import logger

from ..core.config import get_settings
from ..models.document import DocumentChunk

settings = get_settings()


class BaseParser(ABC):
    """文档解析器基类"""
    
    def __init__(self):
        self.logger = logger.bind(name=self.__class__.__name__)
    
    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析该文件"""
        pass
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档内容"""
        pass
    
    def split_text(self, text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
        """文本分段处理"""
        if chunk_size is None:
            chunk_size = settings.chunk_size
        if chunk_overlap is None:
            chunk_overlap = settings.chunk_overlap
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # 如果不是最后一段，尝试在句号、问号、感叹号处分割
            if end < len(text):
                # 寻找最近的句子结束标记
                for i in range(end, max(start, end - 100), -1):
                    if text[i] in '.!?。！？':
                        end = i + 1
                        break
                
                # 如果没有找到句子结束标记，尝试在空格处分割
                if end == start + chunk_size:
                    for i in range(end, max(start, end - 50), -1):
                        if text[i].isspace():
                            end = i + 1
                            break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 计算下一段的起始位置，考虑重叠
            start = max(start + 1, end - chunk_overlap)
        
        return chunks


class PDFParser(BaseParser):
    """PDF文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析PDF文件"""
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str) -> str:
        """解析PDF文档内容"""
        try:
            self.logger.info(f"开始解析PDF文件: {file_path}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                text_content = []
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text.strip())
                        
                        # 记录进度
                        if (page_num + 1) % 10 == 0:
                            self.logger.info(f"PDF解析进度: {page_num + 1}/{total_pages}")
                    
                    except Exception as e:
                        self.logger.warning(f"解析PDF第{page_num + 1}页时出错: {e}")
                        continue
                
                full_text = '\n\n'.join(text_content)
                self.logger.info(f"PDF文件解析完成: {file_path}, 总页数: {total_pages}, 文本长度: {len(full_text)}")
                
                return full_text
                
        except Exception as e:
            self.logger.error(f"解析PDF文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"PDF解析失败: {e}")


class DocxParser(BaseParser):
    """Word文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析Word文件"""
        return file_path.lower().endswith('.docx')
    
    def parse(self, file_path: str) -> str:
        """解析Word文档内容"""
        try:
            self.logger.info(f"开始解析Word文件: {file_path}")
            
            doc = Document(file_path)
            
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_content)
            self.logger.info(f"Word文件解析完成: {file_path}, 段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}, 文本长度: {len(full_text)}")
            
            return full_text
            
        except Exception as e:
            self.logger.error(f"解析Word文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"Word解析失败: {e}")


class MarkdownParser(BaseParser):
    """Markdown文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析Markdown文件"""
        return file_path.lower().endswith(('.md', '.markdown'))
    
    def parse(self, file_path: str) -> str:
        """解析Markdown文档内容"""
        try:
            self.logger.info(f"开始解析Markdown文件: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # 转换为纯文本（移除Markdown标记）
            # 这里使用简单的正则表达式移除Markdown标记
            # 对于更复杂的Markdown，可以考虑使用专门的库
            
            # 移除标题标记
            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
            
            # 移除粗体和斜体标记
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            content = re.sub(r'\*(.*?)\*', r'\1', content)
            
            # 移除代码块标记
            content = re.sub(r'```.*?\n', '', content, flags=re.DOTALL)
            content = re.sub(r'```', '', content)
            
            # 移除行内代码标记
            content = re.sub(r'`(.*?)`', r'\1', content)
            
            # 移除链接标记
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            
            # 移除图片标记
            content = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', content)
            
            # 移除列表标记
            content = re.sub(r'^[\s]*[-*+]\s+', '', content, flags=re.MULTILINE)
            content = re.sub(r'^[\s]*\d+\.\s+', '', content, flags=re.MULTILINE)
            
            # 移除引用标记
            content = re.sub(r'^>\s+', '', content, flags=re.MULTILINE)
            
            # 清理多余的空行
            content = re.sub(r'\n\s*\n', '\n\n', content)
            content = content.strip()
            
            self.logger.info(f"Markdown文件解析完成: {file_path}, 文本长度: {len(content)}")
            
            return content
            
        except Exception as e:
            self.logger.error(f"解析Markdown文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"Markdown解析失败: {e}")


class TxtParser(BaseParser):
    """纯文本文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析文本文件"""
        return file_path.lower().endswith('.txt')
    
    def parse(self, file_path: str) -> str:
        """解析文本文件内容"""
        try:
            self.logger.info(f"开始解析文本文件: {file_path}")
            
            # 尝试不同的编码格式
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise RuntimeError("无法使用支持的编码格式读取文件")
            
            # 清理文本
            content = content.strip()
            
            self.logger.info(f"文本文件解析完成: {file_path}, 文本长度: {len(content)}")
            
            return content
            
        except Exception as e:
            self.logger.error(f"解析文本文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"文本解析失败: {e}")


class DocumentParserService:
    """文档解析服务"""
    
    def __init__(self):
        self.logger = logger.bind(name=self.__class__.__name__)
        self.parsers = [
            PDFParser(),
            DocxParser(),
            MarkdownParser(),
            TxtParser()
        ]
    
    def get_parser(self, file_path: str) -> BaseParser:
        """获取适合的解析器"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        
        raise ValueError(f"不支持的文件格式: {file_path}")
    
    def parse_document(self, file_path: str) -> str:
        """解析文档内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        parser = self.get_parser(file_path)
        return parser.parse(file_path)
    
    def create_chunks(
        self, 
        document_id: str, 
        content: str, 
        chunk_size: int = None, 
        chunk_overlap: int = None
    ) -> List[DocumentChunk]:
        """创建文档片段"""
        parser = BaseParser()  # 使用基类的方法
        
        # 文本分段
        text_chunks = parser.split_text(content, chunk_size, chunk_overlap)
        
        # 创建DocumentChunk对象
        chunks = []
        current_pos = 0
        
        for i, chunk_text in enumerate(text_chunks):
            chunk = DocumentChunk(
                document_id=document_id,
                content=chunk_text,
                chunk_index=i,
                start_char=current_pos,
                end_char=current_pos + len(chunk_text),
                metadata={
                    "chunk_size": len(chunk_text),
                    "chunk_index": i
                }
            )
            chunks.append(chunk)
            current_pos += len(chunk_text)
        
        self.logger.info(f"创建文档片段完成: 文档ID={document_id}, 片段数={len(chunks)}")
        
        return chunks
    
    def parse_and_chunk(
        self, 
        document_id: str, 
        file_path: str, 
        chunk_size: int = None, 
        chunk_overlap: int = None
    ) -> tuple[str, List[DocumentChunk]]:
        """解析文档并创建片段"""
        # 解析文档内容
        content = self.parse_document(file_path)
        
        # 创建片段
        chunks = self.create_chunks(document_id, content, chunk_size, chunk_overlap)
        
        return content, chunks
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return ['.pdf', '.docx', '.md', '.markdown', '.txt']
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """验证文件"""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            validation_result["valid"] = False
            validation_result["errors"].append("文件不存在")
            return validation_result
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > settings.max_file_size:
            validation_result["valid"] = False
            validation_result["errors"].append(f"文件大小超过限制: {file_size} > {settings.max_file_size}")
        
        # 检查文件格式
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.get_supported_formats():
            validation_result["valid"] = False
            validation_result["errors"].append(f"不支持的文件格式: {file_ext}")
        
        # 检查文件是否可读
        try:
            with open(file_path, 'rb') as f:
                f.read(1024)  # 尝试读取前1KB
        except Exception as e:
            validation_result["valid"] = False
            validation_result["errors"].append(f"文件无法读取: {e}")
        
        return validation_result
